from odt import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement

# Creare un nuovo documento Word
doc = Document()
doc.add_heading('Business Rules', 0)

# Aggiungere un'intestazione
section = doc.sections[0]
header = section.header
header_para = header.add_paragraph()
header_para.text = 'Business Rules Document'
header_para.style.font.size = Pt(14)
header_para.style.font.bold = True
header_para.style.font.color.rgb = RGBColor(0, 0, 0)

# Aggiungere la sezione dei vincoli di integrità
doc.add_heading('Vincoli di Integrità', level=1)

vincoli = [
    ("RV1", "Un utente registrato deve fornire nome utente, password, data di nascita, numero di telefono o indirizzo email."),
    ("RV2", "Un utente registrato deve poter chattare, seguire streamer, abbonarsi, ricevere notifiche e valutare contenuti."),
    ("RV3", "Un utente non registrato non deve poter chattare, abbonarsi, ricevere notifiche o valutare contenuti."),
    ("RV4", "Uno streamer deve possedere un canale."),
    ("RV5", "Un canale deve avere una descrizione, un trailer, una lista di social associati e una immagine profilo."),
    ("RV6", "Le live devono avere un titolo, una durata e appartenere a una categoria."),
    ("RV7", "Gli utenti devono poter cercare contenuti tramite hashtag, emoji o categoria."),
    ("RV8", "I contenuti multimediali devono essere gestiti da una piattaforma di video hosting esterna."),
    ("RV9", "Lo streamer deve pagare un corrispettivo di $50 mensili per il servizio di hosting del canale."),
    ("RV10", "Se lo streamer non paga il corrispettivo entro 15 giorni, il canale deve essere sospeso fino al pagamento.")
]

table_vincoli = doc.add_table(rows=1, cols=2)
hdr_cells = table_vincoli.rows[0].cells
hdr_cells[0].text = "Codice"
hdr_cells[1].text = "Descrizione"

# Formattare l'intestazione della tabella
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(12)
    cell.paragraphs[0].alignment = 1  # Centra il testo

for vincolo in vincoli:
    row_cells = table_vincoli.add_row().cells
    row_cells[0].text = vincolo[0]
    row_cells[1].text = vincolo[1]

# Aggiungere la sezione delle derivazioni
doc.add_heading('Derivazioni', level=1)

derivazioni = [
    ("RD1", "Un utente premium si ottiene quando un utente registrato paga un abbonamento."),
    ("RD2", "Il numero di live effettuate da uno streamer si ottiene conteggiando tutte le trasmissioni live effettuate."),
    ("RD3", "Il numero di minuti trasmessi da uno streamer si ottiene sommando la durata di tutte le live e video."),
    ("RD4", "Il numero medio di spettatori di uno streamer si ottiene dividendo il totale degli spettatori per il numero di live."),
    ("RD5", "Il numero di follower di uno streamer si ottiene conteggiando gli utenti registrati che seguono il canale."),
    ("RD6", "Lo status di affiliate si ottiene quando uno streamer trasmette almeno 500 minuti, ha una media di tre spettatori simultanei e almeno 50 follower."),
    ("RD7", "La quantità di bit in un portafoglio si ottiene sommando le transazioni di acquisto di bit effettuate dall'utente.")
]

table_derivazioni = doc.add_table(rows=1, cols=2)
hdr_cells = table_derivazioni.rows[0].cells
hdr_cells[0].text = "Codice"
hdr_cells[1].text = "Descrizione"

# Formattare l'intestazione della tabella
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(12)
    cell.paragraphs[0].alignment = 1  # Centra il testo

for derivazione in derivazioni:
    row_cells = table_derivazioni.add_row().cells
    row_cells[0].text = derivazione[0]
    row_cells[1].text = derivazione[1]

# Aggiungere una sezione per note e contatti
doc.add_heading('Note', level=1)
doc.add_paragraph("Aggiungere qui eventuali note o commenti aggiuntivi.")

doc.add_heading('Contatti', level=1)
doc.add_paragraph("Per ulteriori informazioni, contattare: [Inserire dettagli di contatto]")

# Salvare il documento
file_path = "/mnt/data/Business_Rules.docx"
doc.save(file_path)

file_path
